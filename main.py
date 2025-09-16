# --- Imports ---
import io
import os
import json
import logging # <-- Added for logging
from typing import List, Optional, Any  # Make sure 'Any' is imported from 'typing'
import io
from PIL import Image

# Office file processing
import docx
import openpyxl
import pptx
import pypdf

# FastAPI and related
from fastapi import (
    FastAPI,
    File,
    Form,
    HTTPException,
    UploadFile,
    Depends,
)
from fastapi.middleware.cors import CORSMiddleware
from starlette.responses import StreamingResponse, JSONResponse

# Google and Gemini
from google import genai

# Environment and security
from dotenv import load_dotenv

# Firebase Authentication
from fastapi.security import HTTPBearer, HTTPAuthorizationCredentials
from firebase_admin import auth, credentials, initialize_app

# --- Configuration & Setup ---
load_dotenv()

# --- NEW: Configure Logging ---
logging.basicConfig(level=logging.INFO, format='%(asctime)s - %(name)s - %(levelname)s - %(message)s')
logger = logging.getLogger(__name__)


app = FastAPI(
    title="Assignment Checker API",
    description="An API to modify Office file metadata and grade assignments using AI.",
    version="3.0.1 (with Logging)",
)

# --- CORS Middleware ---
origins = [
    "http://127.0.0.1:5500",
    "http://localhost:5500",
    "http://localhost:3000",
    "https://vip-hw.web.app/"
]

app.add_middleware(
    CORSMiddleware,
    allow_origins=origins,
    allow_credentials=True,
    allow_methods=["*"],
    allow_headers=["*"],
)

# --- Firebase Admin SDK Initialization ---

try:
    # Check for the Railway environment variable first
    firebase_creds_json_str = os.environ.get("FIREBASE_SERVICE_ACCOUNT_JSON")

    if firebase_creds_json_str:
        # If the variable exists (on Railway), load credentials from it
        logger.info("Initializing Firebase Admin SDK from environment variable.")
        firebase_creds = json.loads(firebase_creds_json_str)
        cred = credentials.Certificate(firebase_creds)
    else:
        # If the variable is not set (local environment), use Application Default Credentials.
        # This automatically finds your firebase-service-account.json file if you've
        # set the GOOGLE_APPLICATION_CREDENTIALS environment variable.
        logger.info("Initializing Firebase Admin SDK using Application Default Credentials.")
        cred = credentials.ApplicationDefault()

    # Initialize the app with whichever credential method was successful
    initialize_app(cred)
    logger.info("Firebase Admin SDK initialized successfully.")

except Exception as e:
    logger.error(f"Error initializing Firebase Admin SDK: {e}")

# --- Authentication Dependency ---
token_auth_scheme = HTTPBearer()

async def verify_firebase_token(token: HTTPAuthorizationCredentials = Depends(token_auth_scheme)):
    """A dependency that verifies the Firebase ID token in the Authorization header."""
    logger.info("Attempting to verify Firebase token...")
    if not token:
        logger.warning("Bearer token was not provided.")
        raise HTTPException(status_code=401, detail="Bearer token not provided.")
    try:
        decoded_token = auth.verify_id_token(token.credentials)
        logger.info(f"Token successfully verified for user UID: {decoded_token.get('uid')}")
        return decoded_token
    except Exception as e:
        logger.error(f"Firebase token verification failed: {e}")
        raise HTTPException(status_code=403, detail=f"Invalid authentication credentials: {e}")


# Get Gemini API Key from environment
GOOGLE_API_KEY = os.environ.get("GOOGLE_API_KEY")

if not GOOGLE_API_KEY:
    logger.warning("GOOGLE_API_KEY environment variable not set.")


# --- Helper Function for Text Extraction ---
async def extract_text_from_file(file: UploadFile) -> str:
    """Extracts text content from various file types (.docx, .pdf, .txt)."""
    filename = file.filename or ""
    content = await file.read()
    await file.seek(0)
    text = ""

    try:
        if filename.endswith(".pdf"):
            reader = pypdf.PdfReader(io.BytesIO(content))
            for page in reader.pages:
                text += page.extract_text() or ""
        elif filename.endswith(".docx"):
            doc = docx.Document(io.BytesIO(content))
            text = "\n".join([para.text for para in doc.paragraphs])
        else:
            text = content.decode("utf-8", errors="ignore")
    except Exception as e:
        logger.error(f"Error processing file {filename}: {e}")
        return f"[Error: Could not read content from file '{filename}']\n"

    return text


# --- API Endpoints ---
@app.post(
    "/get-metadata/",
    summary="Get Office File Metadata",
    dependencies=[Depends(verify_firebase_token)],
)
async def get_file_metadata(
    file: UploadFile = File(..., description="The Office file to inspect.")
) -> JSONResponse:
    """Reads an Office file and returns its author and last modified by properties."""
    logger.info(f"POST /get-metadata/ endpoint hit for file: {file.filename}")
    file_content = await file.read()
    file_stream = io.BytesIO(file_content)
    filename = str(file.filename)
    file_extension = os.path.splitext(filename)[1]
    metadata = {"author": "", "last_modified_by": ""}

    try:
        if file_extension == ".docx":
            doc = docx.Document(file_stream)
            props = doc.core_properties
            metadata["author"] = props.author or ""
            metadata["last_modified_by"] = props.last_modified_by or ""
        elif file_extension == ".pptx":
            prs = pptx.Presentation(file_stream)  # type: ignore
            props = prs.core_properties
            metadata["author"] = props.author or ""
            metadata["last_modified_by"] = props.last_modified_by or ""
        elif file_extension == ".xlsx":
            wb = openpyxl.load_workbook(file_stream)
            props = wb.properties
            metadata["author"] = props.creator or "" # type: ignore
            metadata["last_modified_by"] = props.lastModifiedBy or "" # type: ignore
        else:
            raise HTTPException(status_code=400, detail="Unsupported file type.")

        return JSONResponse(content=metadata)
    except Exception as e:
        logger.error(f"Failed to read metadata for {file.filename}: {e}")
        raise HTTPException(
            status_code=500, detail=f"Failed to read metadata: {str(e)}"
        )


@app.post(
    "/process-file/",
    summary="Modify Office File Metadata",
    dependencies=[Depends(verify_firebase_token)],
)
async def modify_file_metadata(
    file: UploadFile = File(
        ..., description="The Office file (.docx, .pptx, .xlsx) to process."
    ),
    author: Optional[str] = Form(None, description="The new author name."),
    last_modified_by: Optional[str] = Form(None, description="The new 'last modified by' name."),
) -> StreamingResponse:
    logger.info(f"POST /process-file/ endpoint hit for file: {file.filename}")
    file_content = await file.read()
    file_stream = io.BytesIO(file_content)
    doc_obj = None
    props = None
    original_filename = str(file.filename)
    file_extension = os.path.splitext(original_filename)[1]

    try:
        if file_extension == ".docx":
            doc_obj = docx.Document(file_stream)
            props = doc_obj.core_properties
        elif file_extension == ".pptx":
            doc_obj = pptx.Presentation(file_stream)  # type: ignore
            props = doc_obj.core_properties
        elif file_extension == ".xlsx":
            doc_obj = openpyxl.load_workbook(file_stream)
            props = doc_obj.properties
        else:
            raise HTTPException(
                status_code=400,
                detail="Unsupported file type. Please upload a .docx, .pptx, or .xlsx file.",
            )

        if author is not None:
            if file_extension == ".xlsx":
                props.creator = author  # type: ignore
            else:
                props.author = author

        if last_modified_by is not None:
            if file_extension == ".xlsx":
                props.lastModifiedBy = last_modified_by  # type: ignore
            else:
                props.last_modified_by = last_modified_by

        output_stream = io.BytesIO()
        doc_obj.save(output_stream)
        output_stream.seek(0)

        return StreamingResponse(
            output_stream,
            media_type=file.content_type,
            headers={"Content-Disposition": f"attachment; filename={original_filename}"},
        )
    except Exception as e:
        logger.error(f"Failed to process file {file.filename}: {e}")
        raise HTTPException(status_code=500, detail=f"Failed to process file: {e}")


@app.post(
    "/check-assignment/",
    summary="Check Assignment with Gemini",
    dependencies=[Depends(verify_firebase_token)],
)
async def check_assignment_with_gemini(
    instructions_text: Optional[str] = Form(None),
    instructions_files: List[UploadFile] = File([]),
    submission_text: Optional[str] = Form(None),
    submission_files: List[UploadFile] = File([]),
) -> dict:
    logger.info("POST /check-assignment/ endpoint hit.")
    if not GOOGLE_API_KEY:
        raise HTTPException(
            status_code=500, detail="GOOGLE_API_KEY is not configured on the server."
        )

    full_instructions = ""
    if instructions_text:
        full_instructions += instructions_text + "\n\n"
    for file in instructions_files:
        full_instructions += await extract_text_from_file(file)

    full_submission = ""
    if submission_text:
        full_submission += submission_text + "\n\n"
    for file in submission_files:
        full_submission += await extract_text_from_file(file)

    if not full_instructions.strip():
        raise HTTPException(
            status_code=400,
            detail="No instructions provided. Please provide either text or files.",
        )
    if not full_submission.strip():
        raise HTTPException(
            status_code=400,
            detail="No submission provided. Please provide either text or files.",
        )

    word_count = len(full_submission.split())

    try:
        client = genai.Client(api_key=GOOGLE_API_KEY)
        
        prompt = f"""
        You are an academic assistant. You will be given instructions for a college homework assignment and a copy of the completed assignment.
        Provide a detailed analysis formatted in Markdown. Your response MUST include the following sections with these exact headings:

        ### Requirements Analysis
        - Use a bulleted list to assess whether each requirement from the instructions was met, partially met, or unmet.

        ### AI Sound Check
        - Provide a brief analysis of the writing style. Does it sound like it was written by a student or does it have hallmarks of AI generation? Explain your reasoning.

        ### Suggestions for Improvement
        - Offer a bulleted list of actionable suggestions for improving the assignment. Focus on areas like deeper analysis, providing more concrete examples, or refining the structure.

        ---
        **ADDITIONAL CONTEXT:**
        * **Accurate Word Count:** {word_count} words. Please use this pre-calculated word count for your analysis, especially when checking for length requirements.
        ---
        
        ASSIGNMENT INSTRUCTIONS:
        {full_instructions}

        ---
        COMPLETED ASSIGNMENT SUBMISSION:
        {full_submission}
        """

        response = await client.aio.models.generate_content(
            model="models/gemini-1.5-flash", contents=prompt
        )

        first_candidate = response.candidates[0] if response.candidates else None
        if (
            first_candidate
            and first_candidate.content
            and first_candidate.finish_reason
            and first_candidate.finish_reason.name == "STOP"
        ):
            content_parts = first_candidate.content.parts
            answer_text = ""
            if content_parts:
                answer_text = "".join(
                    part.text
                    for part in content_parts
                    if hasattr(part, "text") and part.text is not None
                )
            if answer_text:
                return {"analysis": answer_text}
            else:
                raise HTTPException(
                    status_code=500, detail="Gemini returned an empty text response."
                )
        else:
            reason = "N/A"
            if first_candidate and first_candidate.finish_reason:
                reason = first_candidate.finish_reason.name
            detail_msg = (
                f"Gemini request did not complete successfully. Reason: {reason}"
            )
            raise HTTPException(status_code=500, detail=detail_msg)
            
    except Exception as e:
        logger.error(f"Error during Gemini API call: {e}")
        raise HTTPException(
            status_code=500, detail=f"An error occurred with the Gemini API: {str(e)}"
        )

# for working on test questions
@app.post(
    "/solve-question/",
    summary="Solve a Question from Images",
    # dependencies=[Depends(verify_firebase_token)], # Uncomment for auth
)
async def solve_question(
    files: List[UploadFile] = File(..., description="One or more images of the test question.")
) -> dict:
    logger.info(f"POST /solve-question/ endpoint hit with {len(files)} file(s).")
    if not GOOGLE_API_KEY:
        raise HTTPException(
            status_code=500, detail="GOOGLE_API_KEY is not configured on the server."
        )

    prompt = "Given a screenshot of a multiple choice question respond with only the question number (if provided) and the letter option of the correct choice (ABCD, etc) unless explicitly directed to do otherwise."

    model_contents: List[Any] = [prompt]
    for file in files:
        if not file.content_type or not file.content_type.startswith("image/"):
            logger.warning(f"Skipping non-image file: {file.filename}")
            continue
        
        image_bytes = await file.read()
        try:
            img = Image.open(io.BytesIO(image_bytes))
            model_contents.append(img)
        except Exception as e:
            logger.error(f"Could not process image file {file.filename}: {e}")
            raise HTTPException(status_code=400, detail=f"Invalid image file: {file.filename}")

    if len(model_contents) <= 1:
        raise HTTPException(status_code=400, detail="No valid image files were provided.")

    try:
        client = genai.Client(api_key=GOOGLE_API_KEY)
        
        response = await client.aio.models.generate_content(
            model="models/gemini-1.5-flash-latest",
            contents=model_contents
        )

        # --- CORRECTED & RESTRUCTURED LOGIC ---
        first_candidate = response.candidates[0] if response.candidates else None

        # 1. Safely determine the finish reason ONCE.
        finish_reason_name = "UNKNOWN"
        if first_candidate and first_candidate.finish_reason:
            if hasattr(first_candidate.finish_reason, 'name'):
                finish_reason_name = first_candidate.finish_reason.name
            else:
                # Handle the case where the reason is a string or other type
                finish_reason_name = str(first_candidate.finish_reason)

        # 2. Now, check the safely-extracted reason.
        if finish_reason_name == "STOP":
            if response.text:
                return {"answer": response.text.strip()}
            else:
                # This can happen if the response is empty but still "STOP"
                raise HTTPException(
                    status_code=500, detail="Gemini returned an empty text response."
                )
        else:
            # If the reason was not "STOP" (e.g., "SAFETY", "MAX_TOKENS", etc.)
            detail_msg = f"Gemini request did not complete successfully. Reason: {finish_reason_name}"
            logger.warning(detail_msg)
            raise HTTPException(status_code=500, detail=detail_msg)

    except Exception as e:
        logger.error(f"Error during Gemini API call for /solve-question/: {e}")
        raise HTTPException(
            status_code=500, detail=f"An error occurred with the Gemini API: {str(e)}"
        )


@app.get(
    "/",
    summary="Root Endpoint",
    description="A simple root endpoint to confirm the API is running.",
)
async def read_root() -> dict:
    return {
        "status": "ok",
        "message": "Welcome to the Assignment Checker API!",
        "docs_url": "/docs",
    }

